#from PyPDF2 import PdfFileReader, PdfFileWriter
import os, sys
import fitz
from PyQt5.QtWidgets import QApplication, QProgressBar, QWidget, QPushButton, QFileDialog, QLabel, QFrame, QMessageBox, QLineEdit
from PyQt5.QtCore import Qt, QThread, pyqtSignal
#from PyQt5.QtGui import QIntValidator
from docx import Document
import win32com
from win32com.client import Dispatch

class MainUI(QWidget):

    def __init__(self):
        super().__init__()
        self.init_ui()

    # 窗体GUI部分
    def init_ui(self):
        # 设置窗体大小
        self.setGeometry(500, 200, 350, 200)
        # 设置固定大小
        self.setFixedSize(520, 370)
        self.setWindowTitle('docx中参考文献识别')
        # 进度条
        self.pbar = QProgressBar(self)
        # 进图条位置及大小
        self.pbar.setGeometry(20, 120, 470, 10)

        # todo 按钮
        # 选择文件按钮
        self.btn_select_file = QPushButton('选择docx文件', self)
        self.btn_select_file.move(20, 230)
        self.btn_select_file.setFixedSize(125, 30)
        self.btn_select_file.clicked.connect(self.file_dialog)
        # # 输出文件按钮
        # self.btn_output_path = QPushButton('选择输出文件夹', self)
        # self.btn_output_path.move(20, 270)
        # self.btn_output_path.setFixedSize(125, 30)
        # self.btn_output_path.clicked.connect(self.output_dialog)
        # 开始按钮
        self.btn = QPushButton('开始', self)
        # 创建按钮并移动
        self.btn.move(150, 310)
        self.btn.setFixedSize(200, 40)
        # 点击按钮，连接事件函数
        self.btn.clicked.connect(self.btn_action)

        # todo 标签
        # 文件路径标签
        self.lab_select_path = QLabel('', self)
        self.lab_select_path.move(150, 230)
        self.lab_select_path.setFixedSize(320, 30)
        self.lab_select_path.setFrameShape(QFrame.Box)
        self.lab_select_path.setFrameShadow(QFrame.Raised)
        # # 输出标签
        # self.lab_output_path = QLabel('文件路径', self)
        # self.lab_output_path.move(150, 270)
        # self.lab_output_path.setFixedSize(320, 30)
        # self.lab_output_path.setFrameShape(QFrame.Box)
        # self.lab_output_path.setFrameShadow(QFrame.Raised)
        # 说明标签
        content = '说明:\n    选择docx文件位置，程序将自动识别文中的参考文献，并以批注的方式查找参考文献缩写。'
        self.description_lab = QLabel(content, self)
        self.description_lab.move(20, 10)
        self.description_lab.setFixedSize(450, 100)
        self.description_lab.setAlignment(Qt.AlignTop)
        self.description_lab.setFrameShape(QFrame.Box)
        self.description_lab.setFrameShadow(QFrame.Raised)
        # 自动换行
        # self.description_lab.adjustSize()
        self.description_lab.setWordWrap(True)
        #
        # # 提取开始页码
        # self.btn_select_file2 = QPushButton('提取开始页码', self)
        # self.btn_select_file2.move(20, 150)
        # self.btn_select_file2.setFixedSize(100, 30)
        #
        # int_validato = QIntValidator(50, 100, self)  # 实例化整型验证器，并设置范围为50-100
        # self.int_le = QLineEdit(self)  # 整型文本框
        # self.int_le.setValidator(int_validato)  # 设置验证
        # self.int_le.setFixedSize(100, 30)
        # self.int_le.move(120, 150)
        #
        # # 截止页码
        # self.btn_select_file1 = QPushButton('提取截止页码', self)
        # self.btn_select_file1.move(260, 150)
        # self.btn_select_file1.setFixedSize(100, 30)
        #
        # int_validato1 = QIntValidator(50, 100, self)  # 实例化整型验证器，并设置范围为50-100
        # self.int_le1 = QLineEdit(self)  # 整型文本框
        # self.int_le1.setValidator(int_validato1)  # 设置验证
        # self.int_le1.setFixedSize(100, 30)
        # self.int_le1.move(360, 150)

        self.step = 0

        # 显示
        self.show()

    # 按钮点击
    def btn_action(self):
        if self.btn.text() == '完成':
            self.close()
        else:
            docx_path = '{}'.format(self.lab_select_path.text())

            if self.btn.text() == '开始':
                if not os.path.exists(docx_path):
                    QMessageBox.warning(self, '', '请选择docx路径', QMessageBox.Yes)
                else:
                    self.btn.setText('程序进行中')
                    self.downloadThread = downloadThread(docx_path)
                    self.downloadThread.download_proess_signal.connect(self.set_progerss_bar)
                    self.downloadThread.start()

    # 选择输入文件路径
    def file_dialog(self):
        # './'表示当前路径
        path = QFileDialog.getOpenFileName(self, '选取文件', './', 'docx文件(*.docx)')
        print(path)
        # 标签框显示文本路径
        self.lab_select_path.setText(path[0])
        # 自动调整标签框大小
        self.lab_select_path.adjustSize()

    def set_progerss_bar(self, num):
        '''
        设置进图条函数
        :param num: 进度条进度（整数）
        :return:
        '''
        self.step = num
        self.pbar.setValue(self.step)
        if num == 100:
            self.btn.setText('完成')
            QMessageBox.information(self, "提示", "批注完成！")
            return


class downloadThread(QThread):

    download_proess_signal = pyqtSignal(int)  # 创建信号

    #num = 0

    def __init__(self, docx_path):
        super(downloadThread, self).__init__()
        self.docx_path = docx_path

    def check_contain_chinese(self, check_str):  # 判断字符串是否含有中文
        for _char in check_str:
            if '\u4e00' <= _char <= '\u9fa5': return True
        return False

    # def insert_comment_byMsg(self, docx_path, err_ref_list):
    #     word = win32com.client.Dispatch('word.Application')  # 打开word应用程序
    #     # word = DispatchEx('Word.Application') # 启动独立的进程
    #     word.Visible = 0  # 后台运行，不打开程序
    #     word.DisplayAlerts = 0  # 不警告
    #     doc = word.Documents.Open(FileName=docx_path, Encoding='gbk')
    #     if len(err_ref_list) > 0:
    #         for err_ref in err_ref_list:
    #             word.Selection.Find.Execute(err_ref['ref'])
    #             comments_str = err_ref['msg']
    #             doc.Comments.Add(Range=word.Selection.Range, Text=comments_str)
    #     doc.Close()

    def readParagraph(self, docx_path):
        '''
        :param docx_path: docx文档的路径
        :return: 返回参考文献列表
        '''

        self.download_proess_signal.emit(1)

        document = Document(docx_path)
        ref_str = False
        i = 1
        ref_list = []
        num = 0
        for paragraph in document.paragraphs:
            ref = paragraph.text.strip()
            if ref_str is True:
                if ref != '':
                    print('[' + str(i) + ']', ref)
                    ref_list.append(ref)
                i = i + 1

            if ref != '':
                ref = ref.replace(' ', '')
                if ref == '参考文献':
                    ref_str = True
            else:
                ref_str = False
            num = int(i / (len(document.paragraphs) - 1 + 15) * 10)
            if num != 100:
                self.download_proess_signal.emit(num)
        return ref_list

    def merge_author(self, merge_arr, ref_obj):
        '''

        :param merge_arr:
        :param ref_obj:
        :return:
        '''
        for obj in merge_arr:
            if obj['author'] == ref_obj['author']:
                obj['year'] = obj['year'].join(ref_obj['year'])
                obj['ref'] = obj['ref'].join(ref_obj['ref'])

        return merge_arr

    def search_equals_names(self, authors, ref_list, year, ref1):
        '''
        :param authors: ['Minshull T A', ' White R S']
        :param ref_list: 所有的参考文献列表
        :return:
        '''
        merge_arr = []
        for ref in ref_list:
            ref_arr = ref.split('.', 2)
            if len(ref_arr) == 3:
                authors1 = ref_arr[0]
                authors_arr = authors1.split(',')

                if '，' in authors1:
                    authors_arr = authors1.split('，')

                year1 = ref_arr[1].strip()

                if year1 == year and authors_arr == authors:
                    print('同一条数据')
                else:
                    if len(authors) < 3:
                        if len(authors) == len(authors_arr):
                            if authors_arr == authors:
                                print(authors,authors_arr,year,year1,ref1,ref)
                                if len(authors) == 1:
                                    ref_obj = {}
                                    ref_obj['author'] = authors
                                    ref_obj['year'] = year

                    else:
                        if authors_arr[0] == authors[0]:
                            print(authors, authors_arr, year, year1, ref1, ref)

    def strQ2B(self, str):
        """把字符串全角转半角"""
        ss = []
        for s in str:
            rstring = ""
            for uchar in s:
                inside_code = ord(uchar)
                if inside_code == 12288:  # 全角空格直接转换
                    inside_code = 32
                elif (inside_code >= 65281 and inside_code <= 65374):  # 全角字符（除空格）根据关系转化
                    inside_code -= 65248
                rstring += chr(inside_code)
            ss.append(rstring)
        return ''.join(ss)

    def create_ref_abbr(self, ref, err_reflist, _comments_mgs):
        """
        :param ref: 一条参考文献
        :return:
        """
        ref_arr = ref.split('.', 2)
        if len(ref_arr) == 3:
            authors = ref_arr[0].strip()
            authors = self.strQ2B(authors)
            authors_arr = authors.split(',')
            year = ref_arr[1].strip()

            if not self.check_contain_chinese(authors):  # 判断是否为英文作者
                if len(authors_arr) == 1:  # 一位英文作者的情况
                    _ref_authors = self._ref_abbr(ref, authors_arr, year)
                    _comments_mgs.append(_ref_authors)
                elif len(authors_arr) == 2:  # 两位作者的情况
                    print('en :', authors_arr)
                    _ref_authors = self._ref_abbr2(ref, authors_arr, year)
                    _comments_mgs.append(_ref_authors)
                elif len(authors_arr) > 2:  # 多位作者的情况
                    print('en :', authors_arr)
                    _ref_authors = self._ref_abbr3(ref, authors_arr, year)
                    _comments_mgs.append(_ref_authors)
                else:
                    print('--------------错误，无作者')
            else:
                if len(authors_arr) == 1:  # 一位作者的情况
                    print('1:', authors_arr[0])
                    _ref_authors = self._ref_abbr_cn(ref, authors_arr, year)
                    # print(_ref_authors)
                    _comments_mgs.append(_ref_authors)
                elif len(authors_arr) == 2:  # 两位作者的情况
                    print('2:', authors_arr)
                    _ref_authors = self._ref_abbr2_cn(ref, authors_arr, year)
                    # print(_ref_authors)
                    _comments_mgs.append(_ref_authors)
                elif len(authors_arr) > 2:  # 多位作者的情况
                    print('3:', authors_arr)
                    _ref_authors = self._ref_abbr3_cn(ref, authors_arr, year)
                    # print(_ref_authors)
                    _comments_mgs.append(_ref_authors)
                else:
                    print('--------------错误，无作者')
        else:
            err_reflist.append(ref)



    def deal_authors_by_ref(self, ref_list):
        '''
        :param ref_list: 根据参考文献列表获取参考文献缩写
        :return:
        '''
        err_refList = []  # 记录错误符号的参考文献
        _comments_mgs = []
        num = 0
        i = 1
        pre_authros = []
        for ref in ref_list:
            self.create_ref_abbr(ref, err_refList, _comments_mgs)
            # ref_arr = ref.split('.', 2)
            # if len(ref_arr) == 3:
            #     authors = ref_arr[0].strip()
            #     authors = self.strQ2B(authors)
            #     print('authors：', authors)
            #     authors_arr = authors.split(',')
            #
            #     year = ref_arr[1].strip()
            #
            #     #self.search_equals_names(authors_arr, ref_list, year, ref)
            #
            #     if not self.check_contain_chinese(authors):  # 判断是否为英文作者
            #         if len(authors_arr) == 1:  # 一位作者的情况
            #             #print('en :', authors_arr)
            #             _ref_authors = self._ref_abbr(ref, authors_arr, year)
            #             _comments_mgs.append(_ref_authors)
            #         elif len(authors_arr) == 2:  # 两位作者的情况
            #             print('en :', authors_arr)
            #             _ref_authors = self._ref_abbr2(ref, authors_arr, year)
            #             #print(_ref_authors)
            #             _comments_mgs.append(_ref_authors)
            #         elif len(authors_arr) > 2:  # 多位作者的情况
            #             print('en :', authors_arr)
            #             _ref_authors = self._ref_abbr3(ref, authors_arr, year)
            #             #print(_ref_authors)
            #             _comments_mgs.append(_ref_authors)
            #         else:
            #             print('--------------错误，无作者')
            #     else:
            #         if len(authors_arr) == 1:  # 一位作者的情况
            #             print('1:', authors_arr[0])
            #             _ref_authors = self._ref_abbr_cn(ref, authors_arr, year)
            #             #print(_ref_authors)
            #             _comments_mgs.append(_ref_authors)
            #         elif len(authors_arr) == 2:  # 两位作者的情况
            #             print('2:', authors_arr)
            #             _ref_authors = self._ref_abbr2_cn(ref, authors_arr, year)
            #             #print(_ref_authors)
            #             _comments_mgs.append(_ref_authors)
            #         elif len(authors_arr) > 2:  # 多位作者的情况
            #             print('3:', authors_arr)
            #             _ref_authors = self._ref_abbr3_cn(ref, authors_arr, year)
            #             #print(_ref_authors)
            #             _comments_mgs.append(_ref_authors)
            #         else:
            #             print('--------------错误，无作者')
            # else:
            #     print('当前参考文献中的点为全角符号，请注意修改。', ref_arr)
            #     err_ref = {}
            #     err_ref['ref'] = ref
            #     err_ref['msg'] = '当前参考文献中的点为全角符号，请注意修改。'
            #     err_refList.append('err_ref')
        #if err_refList: self.insert_comment_byMsg(err_refList)
            i = i + 1
            num = int(i / (len(ref_list) - 1 + 15) * 5)
            if num != 100:
                self.download_proess_signal.emit(num)

        if len(_comments_mgs) > 0:
            for _author in _comments_mgs:
                authors_names = _author['names']
                print(authors_names)

        print(err_refList)
        print(_comments_mgs)

        return _comments_mgs



    def _ref_abbr(self, ref, authors_arr, year):
        '''
            authors_arr : #['Brown M M'] 一位英文作者
            ref_arr:
        '''
        _ref_authors = {}
        _name = []
        first_name = authors_arr[0].strip().split()[0].strip()
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '（' + year + "）")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '(' + year + ")")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '，' + year)
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + ',' + year)

        _ref_authors['ref'] = ref
        _ref_authors['names'] = _name
        _ref_authors['a_names'] = authors_arr
        _ref_authors['year'] = year


        #print(_ref_authors)
        return _ref_authors

    '''
        authors_arr : #['Brown M'] 一位中文作者  
        ref_arr: 
    '''

    def _ref_abbr_cn(self, ref, authors_arr, year):  #
        print(authors_arr)
        if '，' in authors_arr[0]:
            return self._ref_abbr3_cn(ref, authors_arr[0].split('，'), year)
        else:

            _ref_authors = {}
            _name = []
            _name.append(authors_arr[0].strip().split(' ')[0].strip() + '（' + year + "）")
            _name.append(authors_arr[0].strip().split(' ')[0].strip() + '(' + year + ")")
            _name.append(authors_arr[0].strip().split(' ')[0].strip() + '，' + year)
            _name.append(authors_arr[0].strip().split(' ')[0].strip() + ',' + year)

            _ref_authors['ref'] = ref
            _ref_authors['names'] = _name

            print(_ref_authors)
            return _ref_authors

    '''
        authors_arr : ['Minshull T A', ' White R S'] 两位英文作者  
        ref_arr: 
    '''

    def _ref_abbr2(self, ref, authors_arr, year):  #
        _ref_authors = {}
        _name = []
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '和' + authors_arr[1].strip().split(' ')[
            0].strip() + '（' + year + "）")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '和' + authors_arr[1].strip().split(' ')[
            0].strip() + '(' + year + ")")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + ' and ' + authors_arr[1].strip().split(' ')[
            0].strip() + '，' + year)
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + ' and ' + authors_arr[1].strip().split(' ')[
            0].strip() + ',' + year)

        _ref_authors['ref'] = ref
        _ref_authors['names'] = _name

        print(_ref_authors)
        return _ref_authors

    '''
        authors_arr :  ['高原', ' 郑斯华'] 两位中文作者  
        ref_arr: 
    '''

    def _ref_abbr2_cn(self, ref, authors_arr, year):  #
        _ref_authors = {}
        _name = []
        first_name = ''
        sec_name = ''
        if not self.check_contain_chinese(authors_arr[0].strip()):
            first_name = authors_arr[0].strip().split(' ')[0].strip()
        else:
            first_name = authors_arr[0].strip()

        if not self.check_contain_chinese(authors_arr[1].strip()):
            sec_name = authors_arr[1].strip().split(' ')[0].strip()
        else:
            sec_name = authors_arr[1].strip()

        _name.append(first_name + '和' + sec_name + '（' + year + "）")
        _name.append(first_name + '和' + sec_name + '(' + year + ")")
        _name.append(first_name + '和' + sec_name + '，' + year)
        _name.append(first_name + '和' + sec_name + ',' + year)
        _name.append(first_name + '、' + sec_name + '，' + year)
        _name.append(first_name + '、' + sec_name + ',' + year)

        _ref_authors['ref'] = ref
        _ref_authors['names'] = _name

        print(_ref_authors)
        return _ref_authors

    '''
        authors_arr : ['Minshull T A', ' Muller M R', ' White R S'], ['Minshull T A', ' Muller M R', ' Robinson C J', ' et al'] 多位作者  
        ref_arr: 
    '''

    def _ref_abbr3(self, ref, authors_arr, year):  #
        _ref_authors = {}
        _name = []
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '等' + '（' + year + "）")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '等' + '(' + year + ")")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + ' et al.' + '，' + year)
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + ' et al.' + ',' + year)

        _ref_authors['ref'] = ref
        _ref_authors['names'] = _name

        print(_ref_authors)
        return _ref_authors

    '''
        authors_arr : ['卫小冬', ' 赵明辉', ' 阮爱国等'] 多位中文作者  
        ref_arr: 
    '''

    def _ref_abbr3_cn(self, ref, authors_arr, year):  #
        _ref_authors = {}
        _name = []


        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '等' + '（' + year + "）")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '等' + '(' + year + ")")
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '等' + '，' + year)
        _name.append(authors_arr[0].strip().split(' ')[0].strip() + '等' + ',' + year)

        _ref_authors['ref'] = ref
        _ref_authors['names'] = _name

        print(_ref_authors)
        return _ref_authors



    def run(self):
        self.download_proess_signal.emit(int(1))
        try:
            #word = win32com.client.Dispatch('Word.Application')  # 打开word应用程序
            word = win32com.client.DispatchEx('word.Application') # 启动独立的进程
            word.Visible = 0  # 后台运行，不打开程序
            word.DisplayAlerts = 0  # 不警告
            doc = word.Documents.Open(FileName=self.docx_path, Encoding='gbk')
            _ref_list = self.readParagraph(self.docx_path)
            _authors_msg = self.deal_authors_by_ref(_ref_list)
            # i = 15
            # num = 0
            # if len(_authors_msg) > 0:
            #     for author in _authors_msg:
            #         #print(author)
            #         names = author['names']
            #         #print(type(names))
            #         for author_name in names:
            #             if author_name != '':
            #                 while word.Selection.Find.Execute(author_name):
            #                     doc.Comments.Add(Range=word.Selection.Range, Text=author['ref']) #给选中的文字添加批注
            #                     word.Selection.Range.HighlightColorIndex = 4
            #                     #print('word.Selection.Range = ', word.Selection.Range)
            #
            #             word.Selection.Start = 0
            #             word.Selection.End = 0
            #
            #         i = i + 1
            #         num = int(i / (len(_authors_msg) - 1 + 15) * 100)
            #         if num != 100:
            #             self.download_proess_signal.emit(num)

            doc.Close()
            word.Quit()

        except Exception as e:
            print(e)
        self.download_proess_signal.emit(int(100))


if __name__ == '__main__':

    print(fitz.VersionBind)
    app = QApplication(sys.argv)
    pbar = MainUI()
    sys.exit(app.exec_())