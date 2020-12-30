from docx import Document
import win32com
from win32com.client import Dispatch
import re



def readParagraph(path):
    document = Document(path)
    ref_str = False
    i = 1
    ref_list = []
    for paragraph in document.paragraphs:
        ref = paragraph.text
        if ref_str is True:
            if ref != '':
                print('['+str(i)+']', ref)
                ref_list.append(ref)
            i = i+1

        if ref != '':
            ref = ref.replace(' ', '')
            if ref == '参考文献':
                ref_str = True
        else:
            ref_str = False

    return ref_list


'''
    处理参考文献中的作者
    ref_list: 参考文献数组
'''
def deal_authors_by_ref(ref_list):
    err_refList = [] # 记录错误符号的参考文献
    for ref in ref_list:
        ref_arr = ref.split('.', 2)
        if len(ref_arr) == 3:
            authors = ref_arr[0]
            # print(authors)
            authors_arr = authors.split(',')

            year = ref_arr[1].strip()
            # print(authors_arr)
            if not check_contain_chinese(authors):  # 判断是否为英文作者
                if len(authors_arr) == 1:  # 一位作者的情况
                    print('en :', authors_arr)
                    _ref_authors = _ref_abbr(ref,  authors_arr, year)
                    print(_ref_authors)
                elif len(authors_arr) == 2:  # 两位作者的情况
                    print('en :', authors_arr)
                    _ref_authors = _ref_abbr2(ref, authors_arr, year)
                    print(_ref_authors)
                elif len(authors_arr) > 2:  # 多位作者的情况
                    print('en :', authors_arr)
                    _ref_authors = _ref_abbr3(ref, authors_arr, year)
                    print(_ref_authors)
                else:
                    print('--------------错误，无作者')
            else:
                if len(authors_arr) == 1:  # 一位作者的情况
                    print('1:', authors_arr[0])
                    _ref_authors = _ref_abbr_cn(ref, authors_arr, year)
                    print(_ref_authors)
                elif len(authors_arr) == 2:  # 两位作者的情况
                    print('2:', authors_arr)
                    _ref_authors = _ref_abbr2_cn(ref, authors_arr, year)
                    print(_ref_authors)
                elif len(authors_arr) > 2:  # 多位作者的情况
                    print('3:', authors_arr)
                    _ref_authors = _ref_abbr3_cn(ref, authors_arr, year)
                    print(_ref_authors)
                else:
                    print('--------------错误，无作者')
        else:
            print('当前参考文献中的点为全角符号，请注意修改。')
            err_ref = {}
            err_ref['ref'] = ref
            err_ref['msg'] = '当前参考文献中的点为全角符号，请注意修改。'
            err_refList.append('err_ref')
    #if err_refList: insert_comment_byMsg(err_refList)


'''
    authors_arr : #['Brown M'] 一位英文作者  
    ref_arr: 
'''
def _ref_abbr(ref, authors_arr, year):#
    _ref_authors = {}
    _name = []
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + '（' + year + "）")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + '(' + year + ")")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + '，' + year)
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + ',' + year)

    _ref_authors['ref'] = ref
    _ref_authors['names'] = _name

    print(_ref_authors)
    return _ref_authors

'''
    authors_arr : #['Brown M'] 一位中文作者  
    ref_arr: 
'''
def _ref_abbr_cn(ref, authors_arr, year):#
    _ref_authors = {}
    _name = []
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + '（' + year + "）")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + '(' + year + ")")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + '，' + year)
    _name.append(authors_arr[0].strip().split(' ')[0].strip() + ',' + year)

    _ref_authors['ref'] = ref
    _ref_authors['names'] = _name

    print(_ref_authors)
    return _ref_authors

'''
    authors_arr : ['Minshull T A', ' White R S'] 两位英文作者  
    ref_arr: 
'''
def _ref_abbr2(ref, authors_arr, year):#
    _ref_authors = {}
    _name = []
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'和' + authors_arr[1].strip().split(' ')[0].strip() + '（' + year + "）")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'和' + authors_arr[1].strip().split(' ')[0].strip() + '(' + year + ")")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +' and ' + authors_arr[1].strip().split(' ')[0].strip() + '，' + year)
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +' and ' + authors_arr[1].strip().split(' ')[0].strip() + ',' + year)

    _ref_authors['ref'] = ref
    _ref_authors['names'] = _name

    print(_ref_authors)
    return _ref_authors

'''
    authors_arr :  ['高原', ' 郑斯华'] 两位中文作者  
    ref_arr: 
'''
def _ref_abbr2_cn(ref, authors_arr, year):#
    _ref_authors = {}
    _name = []
    _name.append(authors_arr[0].strip() +'和' + authors_arr[1].strip() + '（' + year + "）")
    _name.append(authors_arr[0].strip() +'和' + authors_arr[1].strip() + '(' + year + ")")
    _name.append(authors_arr[0].strip() +'和' + authors_arr[1].strip() + '，' + year)
    _name.append(authors_arr[0].strip() +'和' + authors_arr[1].strip() + ',' + year)
    _name.append(authors_arr[0].strip() + '、' + authors_arr[1].strip() + '，' + year)
    _name.append(authors_arr[0].strip() + '、' + authors_arr[1].strip() + ',' + year)

    _ref_authors['ref'] = ref
    _ref_authors['names'] = _name

    print(_ref_authors)
    return _ref_authors



'''
    authors_arr : ['Minshull T A', ' Muller M R', ' White R S'], ['Minshull T A', ' Muller M R', ' Robinson C J', ' et al'] 多位作者  
    ref_arr: 
'''
def _ref_abbr3(ref, authors_arr, year):#
    _ref_authors = {}
    _name = []
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'等' + '（' + year + "）")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'等' + '(' + year + ")")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +' et al.' + '，' + year)
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +' et al.' + ',' + year)

    _ref_authors['ref'] = ref
    _ref_authors['names'] = _name

    print(_ref_authors)
    return _ref_authors

'''
    authors_arr : ['卫小冬', ' 赵明辉', ' 阮爱国等'] 多位中文作者  
    ref_arr: 
'''
def _ref_abbr3_cn(ref, authors_arr, year):#
    _ref_authors = {}
    _name = []
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'等' + '（' + year + "）")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'等' + '(' + year + ")")
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'等' + '，' + year)
    _name.append(authors_arr[0].strip().split(' ')[0].strip() +'等' + ',' + year)

    _ref_authors['ref'] = ref
    _ref_authors['names'] = _name

    print(_ref_authors)
    return _ref_authors

'''
    调用win32com的系统服务word，向指定位置插入批注。
    docx_path: word的路径
    ref_list: 插入多个批注的数组
'''

def insert_comment(docx_path, ref_list):
    word = win32com.client.Dispatch('word.Application') # 打开word应用程序
    # word = DispatchEx('Word.Application') # 启动独立的进程
    word.Visible = 0 #后台运行，不打开程序
    word.DisplayAlerts = 0  # 不警告
    doc = word.Documents.Open(FileName = docx_path, Encoding = 'gbk')
    if len(ref_list) > 0:
        for where_str in ref_list:
            word.Selection.Find.Execute(where_str)
            comments_str = where_str
            # s = word.Selection
            # s.Start = 0
            # s.End = 0
            doc.Comments.Add(Range = word.Selection.Range, Text = comments_str)

    # content = doc.Range(doc.Content.Start, doc.Content.End)
    # content = doc.Range()
    # print('----------------')
    # print('段落数: ', doc.Paragraphs.count)
    # # 利用下标遍历段落
    # for i in range(len(doc.Paragraphs)):
    #     para = doc.Paragraphs[i]
    #     print(i, para.Range.text)
    #     print('-------------------------')
        # 直接遍历段落
    #for para in doc.paragraphs:
    #    print(para.Range.text)
        # print para #只能用于文档内容全英文的情况
    doc.Close()  # 关闭word文档
    # word.Quit #关闭word程序


def insert_comment_byMsg(docx_path, err_ref_list):
    word = win32com.client.Dispatch('word.Application') # 打开word应用程序
    # word = DispatchEx('Word.Application') # 启动独立的进程
    word.Visible = 0 #后台运行，不打开程序
    word.DisplayAlerts = 0  # 不警告
    doc = word.Documents.Open(FileName = docx_path, Encoding = 'gbk')
    if len(err_ref_list) > 0:
        for err_ref in err_ref_list:
            word.Selection.Find.Execute(err_ref['ref'])
            comments_str = err_ref['msg']
            doc.Comments.Add(Range = word.Selection.Range, Text = comments_str)
    doc.Close()


def check_contain_chinese(check_str): #判断是否含有中文
    for _char in check_str:
        if '\u4e00' <= _char <= '\u9fa5': return True
    return False

if __name__ == '__main__':
    docx_path = 'C:/Users/dell/Documents/Tencent Files/495147399/FileRecv/海底地震勘测理论基础与实用技术 7-#.docx'
    #print(check_contain_chinese('sfsd  sf'))
    ref_list = readParagraph(docx_path)

    if len(ref_list) > 0:
        deal_authors_by_ref(ref_list)
        #insert_comment(docx_path, ref_list)
    a = []
    if a: print(' null')